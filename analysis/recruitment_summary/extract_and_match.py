import json
import re
import sqlite3
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "docs" / "招聘信息汇总.docx"
DB_PATH = ROOT / "db.sqlite3"
OUTPUT_PATH = Path(__file__).with_name("data.json")

OWNER_PATTERN = re.compile(r"招聘负责人[：:]\s*([^）\n]+)")
PARENTHETICAL_OWNER_PATTERN = re.compile(r"[（(]\s*招聘负责人[：:][^）)]*[）)]")
NUMBER_PREFIX_PATTERN = re.compile(r"^\s*\d+\s*[、.．]\s*")
STATUS_PATTERN = re.compile(r"\s*(暂不招聘|暂停招聘)\s*$")


def split_aliases(title):
    aliases = []
    buffer = []
    depth = 0
    for character in title:
        if character in "（(":
            depth += 1
        elif character in "）)" and depth:
            depth -= 1
        if character == "/" and depth == 0:
            alias = "".join(buffer).strip()
            if alias:
                aliases.append(alias)
            buffer = []
        else:
            buffer.append(character)
    alias = "".join(buffer).strip()
    if alias:
        aliases.append(alias)
    return aliases


def parse_owners(text):
    match = OWNER_PATTERN.search(text)
    if not match:
        return []
    owner_text = match.group(1).strip().rstrip("）)").strip()
    return [
        owner.strip()
        for owner in re.split(r"[、,，/]", owner_text)
        if owner.strip()
    ]


def clean_job_title(text):
    title = NUMBER_PREFIX_PATTERN.sub("", text).strip()
    title = PARENTHETICAL_OWNER_PATTERN.sub("", title).strip()
    title = OWNER_PATTERN.sub("", title).strip()
    title = STATUS_PATTERN.sub("", title).strip()
    return title


def extract_document_jobs():
    document = Document(DOCX_PATH)
    jobs = []
    department = ""
    department_owners = []

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        if not text:
            continue

        if style == "Heading 1":
            if text.startswith("（招聘负责人"):
                department_owners = parse_owners(text)
                continue
            if index >= 19:
                department = re.sub(r"^[一二三四五六七八九十]+、", "", text)
                department = OWNER_PATTERN.sub("", department).strip()
                department_owners = parse_owners(text)
            continue

        if style != "Heading 2":
            continue

        next_heading_index = len(document.paragraphs)
        for next_index in range(index + 1, len(document.paragraphs)):
            next_style = (
                document.paragraphs[next_index].style.name
                if document.paragraphs[next_index].style
                else ""
            )
            if next_style in {"Heading 1", "Heading 2"}:
                next_heading_index = next_index
                break

        block = [
            document.paragraphs[block_index].text.strip()
            for block_index in range(index + 1, next_heading_index)
            if document.paragraphs[block_index].text.strip()
        ]
        explicit_owners = parse_owners(text)
        if not explicit_owners:
            for block_text in block[:3]:
                explicit_owners = parse_owners(block_text)
                if explicit_owners:
                    break

        status = "暂停招聘" if STATUS_PATTERN.search(text) else "正常"
        title = clean_job_title(text)
        content = "\n".join(
            line for line in block if not OWNER_PATTERN.search(line)
        )
        jobs.append(
            {
                "department": department,
                "title": title,
                "aliases": split_aliases(title),
                "owners": explicit_owners or department_owners,
                "owner_source": "岗位" if explicit_owners else "部门",
                "status": status,
                "content": content,
                "paragraph_index": index,
            }
        )

    return jobs


def load_beisen_jobs():
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    jobs = []
    for row in connection.execute(
        """
        SELECT name, beisen_position_id, status, status_source, source_payload
        FROM recruitment_position
        ORDER BY name
        """
    ):
        payload = json.loads(row["source_payload"] or "{}")
        hr_owner = payload.get("hrDutyUserLite") or {}
        jobs.append(
            {
                "title": row["name"],
                "beisen_position_id": row["beisen_position_id"],
                "status": row["status"],
                "status_source": row["status_source"],
                "hr_owner": hr_owner.get("name", ""),
                "hr_email": hr_owner.get("email", ""),
                "share_user_ids": payload.get("shareUserIds") or [],
                "content": "\n".join(
                    value
                    for value in [payload.get("duty", ""), payload.get("require", "")]
                    if value
                ),
            }
        )
    connection.close()
    return jobs


def normalize_title(value):
    value = value.lower()
    replacements = {
        "amazon": "亚马逊",
        "tiktok": "tk",
        "aigc": "ai",
        "（": "(",
        "）": ")",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)
    value = re.sub(r"\((双休|出海赛道|跨境电商方向|应届生可投|英语方向|面向海外)\)", "", value)
    value = re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", value)
    return value


def normalize_content(value):
    value = value.lower()
    value = re.sub(r"岗位职责|任职要求|职位要求", "", value)
    value = re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", value)
    return value


def ngrams(value, size=3):
    if len(value) < size:
        return {value} if value else set()
    return {value[index : index + size] for index in range(len(value) - size + 1)}


def jaccard(left, right):
    left_set = ngrams(normalize_content(left))
    right_set = ngrams(normalize_content(right))
    if not left_set or not right_set:
        return 0.0
    return len(left_set & right_set) / len(left_set | right_set)


def title_similarity(left, aliases):
    normalized_left = normalize_title(left)
    scores = []
    for alias in aliases:
        normalized_alias = normalize_title(alias)
        if not normalized_alias:
            continue
        if normalized_left == normalized_alias:
            scores.append(1.0)
            continue
        if normalized_left in normalized_alias or normalized_alias in normalized_left:
            containment = min(len(normalized_left), len(normalized_alias)) / max(
                len(normalized_left), len(normalized_alias)
            )
            scores.append(max(0.82, containment))
            continue
        scores.append(SequenceMatcher(None, normalized_left, normalized_alias).ratio())
    return max(scores, default=0.0)


def build_candidates(beisen_jobs, document_jobs):
    matches = []
    for beisen_job in beisen_jobs:
        candidates = []
        for document_job in document_jobs:
            title_score = title_similarity(
                beisen_job["title"], document_job["aliases"]
            )
            content_score = jaccard(
                beisen_job["content"], document_job["content"]
            )
            combined_score = round(0.45 * title_score + 0.55 * content_score, 4)
            candidates.append(
                {
                    "document_title": document_job["title"],
                    "document_owners": document_job["owners"],
                    "document_status": document_job["status"],
                    "title_score": round(title_score, 4),
                    "content_score": round(content_score, 4),
                    "combined_score": combined_score,
                }
            )
        candidates.sort(key=lambda item: item["combined_score"], reverse=True)
        matches.append(
            {
                **beisen_job,
                "candidates": candidates[:5],
            }
        )
    return matches


document_jobs = extract_document_jobs()
beisen_jobs = load_beisen_jobs()
matches = build_candidates(beisen_jobs, document_jobs)
OUTPUT_PATH.write_text(
    json.dumps(
        {
            "document_jobs": document_jobs,
            "beisen_jobs": beisen_jobs,
            "matches": matches,
        },
        ensure_ascii=False,
        indent=2,
    ),
    encoding="utf-8",
)
print(f"document_jobs={len(document_jobs)} beisen_jobs={len(beisen_jobs)}")
for match in matches:
    best = match["candidates"][0]
    print(
        f"{match['title']} -> {best['document_title']} "
        f"title={best['title_score']:.3f} jd={best['content_score']:.3f} "
        f"combined={best['combined_score']:.3f}"
    )
