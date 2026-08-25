import hashlib
import re
from io import BytesIO

from django.db import transaction
from django.db.models import Max
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook

from recruitment.models import (
    DocumentPosition,
    Position,
    PositionConfiguration,
    ReferenceDocument,
)
from recruitment.services.common import record_audit
from recruitment.services.position_matching import (
    apply_match_suggestion,
    normalize_position_title,
)
from reviews.models import PositionReviewer, Reviewer


def clean_position_title(title):
    title = str(title or "").strip()
    title = re.sub(r"[（(]\s*(?:招聘)?负责人[：:][^）)]*[）)]?", "", title)
    title = re.sub(r"[（(]\s*(?:暂不招聘|暂停招聘|待招)\s*[）)]?", "", title)
    title = re.sub(r"\s*(?:招聘)?负责人[：:].*$", "", title)
    title = re.sub(r"\s*(?:暂不招聘|暂停招聘|待招)\s*$", "", title)
    title = re.sub(r"^\d+\s*[、.．]\s*", "", title)
    return title.strip()


def split_values(value):
    return [
        item.strip()
        for item in re.split(r"[；;、,，\n/|]+", str(value or ""))
        if item and item.strip()
    ]


def _table_records(document):
    records = []
    for table in document.tables:
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        header_index = next(
            (
                index
                for index, row in enumerate(rows)
                if any("岗位" in cell or "职位" in cell for cell in row)
            ),
            None,
        )
        if header_index is None:
            continue
        headers = rows[header_index]
        for row in rows[header_index + 1 :]:
            values = {
                headers[index].strip(): row[index].strip()
                for index in range(min(len(headers), len(row)))
            }
            title = next(
                (
                    value
                    for key, value in values.items()
                    if ("岗位" in key or "职位" in key) and value
                ),
                "",
            )
            title = clean_position_title(title)
            if not title:
                continue
            jd_parts = [
                value
                for key, value in values.items()
                if any(word in key for word in ("职责", "要求", "说明", "JD"))
                and value
            ]
            aliases = [clean_position_title(a) for a in split_values(title)]
            records.append(
                {
                    "title": title,
                    "aliases": [a for a in aliases if a],
                    "jd": "\n".join(jd_parts),
                    "source_section": "表格",
                    "metadata": values,
                }
            )
    return records


def parse_job_summary_docx(content):
    document = Document(BytesIO(content))
    records = _table_records(document)
    if records:
        return records

    current = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        title_match = re.match(r"^(?:岗位名称|职位名称)[:：]\s*(.+)$", text)
        is_heading = style_name.lower().startswith("heading")
        if title_match or is_heading:
            raw_title = (title_match.group(1) if title_match else text).strip()
            title = clean_position_title(raw_title)
            if title and len(title) <= 80 and not any(
                word in title
                for word in ("招聘信息", "岗位职责", "任职要求", "注意事项", "筛选注意")
            ):
                aliases = [clean_position_title(a) for a in split_values(title)]
                current = {
                    "title": title,
                    "aliases": [a for a in aliases if a],
                    "jd_lines": [],
                    "source_section": style_name or "正文",
                    "metadata": {},
                }
                records.append(current)
                continue
        if current:
            current["jd_lines"].append(text)
    parsed = []
    for record in records:
        jd = record.get("jd") or "\n".join(record.pop("jd_lines", []))
        if jd:
            record["jd"] = jd
            parsed.append(record)
    if not parsed:
        raise ValueError("未识别到岗位名称和岗位说明，请检查文档结构。")
    return parsed


def parse_reviewer_mapping_xlsx(content):
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    records = []
    for worksheet in workbook.worksheets:
        rows = list(worksheet.iter_rows(values_only=True))
        for header_index, row in enumerate(rows):
            headers = [str(value or "").strip() for value in row]
            pos_columns = [
                index
                for index, value in enumerate(headers)
                if ("岗位" in value or "职位" in value)
                and not any(k in value for k in ("ID", "说明", "JD", "状态", "相似度"))
            ]
            owner_column = next(
                (index for index, value in enumerate(headers) if "负责人" in value and "邮箱" not in value and "HR" not in value),
                None,
            )
            if owner_column is None:
                owner_column = next(
                    (index for index, value in enumerate(headers) if "负责人" in value and "邮箱" not in value),
                    None,
                )
            email_column = next(
                (index for index, value in enumerate(headers) if "邮箱" in value and "HR" not in value),
                None,
            )
            if email_column is None:
                email_column = next(
                    (index for index, value in enumerate(headers) if "邮箱" in value),
                    None,
                )
            if not pos_columns or owner_column is None or email_column is None:
                continue
            title_column = pos_columns[0]
            for data_row in rows[header_index + 1 :]:
                if max(title_column, owner_column) >= len(data_row):
                    continue
                title = clean_position_title(data_row[title_column])
                if not title or title.startswith("未找到"):
                    continue
                aliases = []
                for col_idx in pos_columns:
                    if col_idx < len(data_row) and data_row[col_idx]:
                        val = clean_position_title(data_row[col_idx])
                        if val and not val.startswith("未找到"):
                            aliases.extend([clean_position_title(x) for x in split_values(val)])
                aliases = list(dict.fromkeys([a for a in aliases if a]))

                raw_owners = str(data_row[owner_column] or "")
                raw_emails = str(data_row[email_column] or "") if email_column < len(data_row) else ""
                names = [n.strip() for n in re.split(r"[、,，；;\s\n]+", raw_owners) if n.strip()]
                emails = [e.strip() for e in re.split(r"[、,，；;\s\n]+", raw_emails) if e.strip() and "@" in e]
                reviewers = []
                for index, name in enumerate(names):
                    email = emails[index] if index < len(emails) else (emails[0] if len(emails) == 1 else "")
                    if name not in {"待确认", "-", "无"} and email and "@" in email:
                        reviewers.append({"name": name, "email": email})
                records.append(
                    {
                        "title": title,
                        "aliases": aliases,
                        "jd": "",
                        "source_section": worksheet.title,
                        "metadata": {"reviewers": reviewers},
                    }
                )
            break
    if not records:
        raise ValueError("未识别到“岗位、负责人、邮箱”列，请检查表格结构。")
    return records


def parse_reference_content(document_type, content):
    if document_type == ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX:
        return parse_job_summary_docx(content)
    if document_type == ReferenceDocument.DocumentType.REVIEWER_MAPPING_XLSX:
        return parse_reviewer_mapping_xlsx(content)
    raise ValueError("不支持的参考资料类型。")


@transaction.atomic
def create_reference_document(uploaded_file, document_type, actor):
    content = uploaded_file.read()
    uploaded_file.seek(0)
    version = (
        ReferenceDocument.objects.filter(document_type=document_type).aggregate(
            value=Max("version")
        )["value"]
        or 0
    ) + 1
    reference = ReferenceDocument.objects.create(
        name=uploaded_file.name,
        document_type=document_type,
        file=uploaded_file,
        content_hash=hashlib.sha256(content).hexdigest(),
        version=version,
        uploaded_by=actor,
    )
    try:
        records = parse_reference_content(document_type, content)
        DocumentPosition.objects.bulk_create(
            [
                DocumentPosition(
                    reference_document=reference,
                    title=record["title"],
                    normalized_title=normalize_position_title(record["title"]),
                    aliases=record.get("aliases", []),
                    jd=record.get("jd", ""),
                    source_section=record.get("source_section", ""),
                    metadata=record.get("metadata", {}),
                )
                for record in records
            ]
        )
    except Exception as exc:
        reference.status = ReferenceDocument.Status.PARSE_FAILED
        reference.parse_error = str(exc)
        reference.save(update_fields=["status", "parse_error"])
    record_audit(actor, "reference_document.upload", reference)
    return reference


def _mapping_for_document_position(document_position, position=None):
    targets = set()
    if document_position:
        clean_title = clean_position_title(document_position.title)
        targets.add(normalize_position_title(clean_title))
        for alias in document_position.aliases or []:
            clean_alias = clean_position_title(alias)
            targets.add(normalize_position_title(clean_alias))
    if position:
        targets.add(normalize_position_title(position.name))
        for alias in split_values(position.name):
            targets.add(normalize_position_title(alias))

    if not targets:
        return None

    mappings = list(
        DocumentPosition.objects.filter(
            reference_document__document_type=ReferenceDocument.DocumentType.REVIEWER_MAPPING_XLSX,
            reference_document__status=ReferenceDocument.Status.ACTIVE,
            is_active=True,
        )
    )
    # 1. Exact match
    for mapping in mappings:
        mapping_names = {
            normalize_position_title(clean_position_title(mapping.title)),
            *(
                normalize_position_title(clean_position_title(alias))
                for alias in mapping.aliases or []
            ),
        }
        if targets & mapping_names:
            return mapping

    # 2. Normalized token / substring match (e.g. 'tiktok' in 'tiktok运营')
    for mapping in mappings:
        mapping_names = [
            normalize_position_title(clean_position_title(mapping.title)),
            *(
                normalize_position_title(clean_position_title(alias))
                for alias in mapping.aliases or []
            ),
        ]
        for t in targets:
            if not t:
                continue
            for m in mapping_names:
                if not m:
                    continue
                if (t in m or m in t) and len(min(t, m)) >= 3:
                    return mapping
    return None


def apply_document_reviewers(configuration, actor):
    mapping = _mapping_for_document_position(
        configuration.document_position,
        position=configuration.position,
    )
    if not mapping:
        return 0
    created = 0
    for item in mapping.metadata.get("reviewers", []):
        reviewer, _ = Reviewer.objects.get_or_create(
            name=item["name"],
            email=item["email"],
        )
        link, was_created = PositionReviewer.objects.get_or_create(
            position=configuration.position,
            reviewer=reviewer,
            defaults={
                "source_type": PositionReviewer.SourceType.DOCUMENT,
                "source_document_position": configuration.document_position,
                "configured_by": actor,
            },
        )
        if was_created:
            record_audit(
                actor,
                "position_reviewer.auto_add",
                link,
                {
                    "document_position_id": configuration.document_position_id,
                    "reviewer_id": reviewer.pk,
                },
            )
        created += int(was_created)
    return created


@transaction.atomic
def publish_reference_document(reference, actor):
    ReferenceDocument.objects.filter(
        document_type=reference.document_type,
        status=ReferenceDocument.Status.ACTIVE,
    ).exclude(pk=reference.pk).update(status=ReferenceDocument.Status.ARCHIVED)
    reference.status = ReferenceDocument.Status.ACTIVE
    reference.published_by = actor
    reference.published_at = timezone.now()
    reference.save(update_fields=["status", "published_by", "published_at"])
    record_audit(actor, "reference_document.publish", reference)

    if reference.document_type == ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX:
        document_positions = list(reference.positions.filter(is_active=True))
        for position in Position.objects.all():
            apply_match_suggestion(
                position,
                document_positions=document_positions,
                force=True,
            )
    else:
        for configuration in PositionConfiguration.objects.filter(
            match_status=PositionConfiguration.MatchStatus.CONFIRMED
        ).select_related("position", "document_position"):
            apply_document_reviewers(configuration, actor)
    return reference
