from io import BytesIO
import json
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch
from uuid import uuid4

import httpx
from django.conf import settings
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management import call_command
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document
from openpyxl import Workbook
from openai import APIConnectionError

from accounts.models import User
from analysis.models import PositionRuleVersion, RuleGenerationOperation
from analysis.services.jobs import AnalysisJobError, create_analysis_job
from analysis.services.rules import (
    RuleDraftError,
    create_generated_rule,
    create_initial_published_rule,
)
from recruitment.models import (
    AuditEvent,
    DocumentPosition,
    Position,
    PositionConfiguration,
    PositionJdDecision,
    ReferenceDocument,
)
from recruitment.services.configuration import (
    confirm_jd,
    configuration_state,
    ensure_position_configuration,
)
from recruitment.services.position_matching import match_position
from recruitment.services.reference_import import (
    apply_document_reviewers,
    parse_job_summary_docx,
    parse_reviewer_mapping_xlsx,
    publish_reference_document,
)
from reviews.models import PositionReviewer, Reviewer
from reviews.services import ReviewError, create_review_batch
from tests.fakes import FakeRuleGateway
from tests.test_workflows import WorkflowFixtureMixin


class PositionConfigurationTests(WorkflowFixtureMixin, TestCase):
    def setUp(self):
        super().setUp()
        self.position = Position.objects.create(
            beisen_position_id="CONFIG-1",
            name="AI算法工程师",
            source_jd="负责模型研发与上线。",
        )
        self.configuration = ensure_position_configuration(self.position)

    def create_document_position(self, title="AI工程师", aliases=None, jd="文档岗位说明"):
        document = ReferenceDocument.objects.create(
            name="招聘汇总",
            document_type=ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX,
            content_hash="hash",
            version=ReferenceDocument.objects.count() + 1,
            status=ReferenceDocument.Status.ACTIVE,
            uploaded_by=self.hr,
        )
        return DocumentPosition.objects.create(
            reference_document=document,
            title=title,
            normalized_title=title.lower(),
            aliases=aliases or [],
            jd=jd,
        )

    def publish_current_rule(self):
        decision = self.position.jd_decisions.get(is_current=True)
        rule = PositionRuleVersion.objects.create(
            position=self.position,
            version=1,
            evaluation_jd=decision.confirmed_jd,
            source_jd_snapshot=decision.confirmed_jd,
            jd_decision=decision,
            hard_requirements=[],
            dimensions=[{"name": "相关经验", "weight": 100}],
            rating_thresholds={"priority": 80, "review": 60},
            created_by=self.hr,
        )
        rule.publish(self.hr)
        return rule

    def add_reviewer(self):
        reviewer = Reviewer.objects.create(name="负责人", email="owner@example.com")
        return PositionReviewer.objects.create(
            position=self.position,
            reviewer=reviewer,
            source_type=PositionReviewer.SourceType.MANUAL,
            configured_by=self.hr,
        )

    def test_new_position_configuration_starts_pending_jd(self):
        state = configuration_state(self.position)
        self.assertEqual(state.code, "pending_jd")
        self.assertIn("确认岗位说明", state.blockers)

    def test_exact_alias_and_ambiguous_matching(self):
        exact = self.create_document_position("AI工程师", ["AI算法工程师"])
        result = match_position(self.position)
        self.assertEqual(result.status, PositionConfiguration.MatchStatus.SUGGESTED)
        self.assertEqual(result.document_position, exact)

        self.create_document_position("AI算法", ["AI算法工程师"])
        result = match_position(self.position)
        self.assertTrue(result.ambiguous)
        self.assertIsNone(result.document_position)

    def test_all_jd_decisions_create_versions_and_audits(self):
        document_position = self.create_document_position()
        self.configuration.document_position = document_position
        self.configuration.match_status = PositionConfiguration.MatchStatus.CONFIRMED
        self.configuration.save()
        first = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        second = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.MANUAL,
            self.hr,
            confirmed_jd="人工确认后的岗位说明",
        )
        third = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.MERGED,
            self.hr,
            confirmed_jd="合并后的岗位说明",
        )
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual([first.version, second.version, third.version], [1, 2, 3])
        self.assertFalse(first.is_current)
        self.assertFalse(second.is_current)
        self.assertTrue(third.is_current)
        self.assertEqual(self.position.refresh_from_db() or self.position.evaluation_jd, "合并后的岗位说明")
        self.assertEqual(
            AuditEvent.objects.filter(action="position_jd.confirm").count(),
            3,
        )

    def test_confirming_unchanged_jd_does_not_create_new_version(self):
        first = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )

        second = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )

        self.assertEqual(first.pk, second.pk)
        self.assertEqual(second.version, 1)
        self.assertTrue(second.is_current)
        self.assertEqual(
            self.position.jd_decisions.count(),
            1,
        )

    def test_configuration_history_collapses_old_versions(self):
        for index in range(9):
            confirm_jd(
                self.position,
                PositionJdDecision.DecisionType.MANUAL,
                self.hr,
                confirmed_jd=f"第 {index + 1} 版岗位说明",
            )

        response = self.authenticated_client(self.hr).get(
            reverse(
                "recruitment:configuration_detail",
                args=[self.position.pk],
            )
        )

        self.assertContains(response, "岗位说明版本")
        self.assertContains(response, 'class="history-scroll"')
        self.assertContains(response, "V9")

    def test_rule_requires_current_decision_and_publish_does_not_change_jd(self):
        with self.assertRaises(RuleDraftError):
            create_generated_rule(self.position, self.hr, FakeRuleGateway())

        decision = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        rule = create_generated_rule(self.position, self.hr, FakeRuleGateway())
        self.assertEqual(rule.jd_decision, decision)
        self.assertEqual(rule.evaluation_jd, decision.confirmed_jd)
        self.position.evaluation_jd = "确认后的内容"
        self.position.save(update_fields=["evaluation_jd"])
        rule.publish(self.hr)
        self.position.refresh_from_db()
        self.assertEqual(self.position.evaluation_jd, "确认后的内容")

    def test_v2_archives_v1_and_update_state_allows_old_rule(self):
        confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        first = self.publish_current_rule()
        self.add_reviewer()
        self.assertEqual(configuration_state(self.position).code, "ready")

        self.position.source_jd = "北森更新后的岗位说明"
        self.position.save(update_fields=["source_jd"])
        state = configuration_state(self.position)
        self.assertEqual(state.code, "update_required")
        self.assertTrue(state.can_run)

        confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        second = PositionRuleVersion.objects.create(
            position=self.position,
            version=2,
            evaluation_jd=self.position.evaluation_jd,
            source_jd_snapshot=self.position.evaluation_jd,
            jd_decision=self.position.jd_decisions.get(is_current=True),
            dimensions=[{"name": "能力", "weight": 100}],
            created_by=self.hr,
        )
        second.publish(self.hr)
        first.refresh_from_db()
        self.assertEqual(first.status, PositionRuleVersion.Status.ARCHIVED)
        self.assertEqual(configuration_state(self.position).code, "ready")

    def test_analysis_rejected_until_configuration_complete(self):
        application = self.create_application(
            applicant_id="CONFIG-APP",
            position=self.position,
        )
        with self.assertRaisesMessage(AnalysisJobError, "确认岗位说明"):
            create_analysis_job(self.position, [application.pk], self.hr)
        confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        self.publish_current_rule()
        job = create_analysis_job(self.position, [application.pk], self.hr)
        self.assertEqual(job.total_count, 1)
        state = configuration_state(self.position)
        self.assertTrue(state.can_analyze)
        self.assertFalse(state.can_review)

    def test_initial_rule_v0_uses_beisen_jd_and_is_immediately_analyzable(self):
        rule, created = create_initial_published_rule(
            self.position,
            self.hr,
            FakeRuleGateway(),
        )

        self.assertTrue(created)
        self.assertEqual(rule.version, 0)
        self.assertEqual(rule.status, PositionRuleVersion.Status.PUBLISHED)
        self.assertEqual(rule.evaluation_jd, self.position.source_jd)
        decision = self.position.jd_decisions.get(is_current=True)
        self.assertEqual(
            decision.decision_type,
            PositionJdDecision.DecisionType.BEISEN,
        )
        state = configuration_state(self.position)
        self.assertTrue(state.can_analyze)
        self.assertFalse(state.can_review)

    def test_initial_reference_document_can_be_seeded_idempotently(self):
        call_command("seed_initial_reference", actor=self.hr.username)
        reference = ReferenceDocument.objects.get(
            document_type=ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX,
            status=ReferenceDocument.Status.ACTIVE,
        )
        position_count = reference.positions.count()
        self.assertGreater(position_count, 0)

        call_command("seed_initial_reference", actor=self.hr.username)

        self.assertEqual(
            ReferenceDocument.objects.filter(
                document_type=ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX,
            ).count(),
            1,
        )
        self.assertEqual(reference.positions.count(), position_count)

    def test_hr_can_access_configuration_and_rules_but_not_admin_pages(self):
        client = self.authenticated_client(self.hr)
        self.assertEqual(
            client.get(reverse("recruitment:configuration_list")).status_code,
            200,
        )
        self.assertEqual(client.get(reverse("analysis:rule_list")).status_code, 200)
        response = client.get(reverse("accounts:user_list"))
        self.assertEqual(response.status_code, 302)

    def test_match_and_jd_confirmation_support_inline_feedback(self):
        document_position = self.create_document_position()
        client = self.authenticated_client(self.hr)
        response = client.post(
            reverse(
                "recruitment:configuration_confirm_match",
                args=[self.position.pk],
            ),
            {"document_position": document_position.pk},
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        self.assertEqual(response.json()["state"]["code"], "pending_jd")

        response = client.post(
            reverse(
                "recruitment:configuration_confirm_jd",
                args=[self.position.pk],
            ),
            {
                "decision_type": PositionJdDecision.DecisionType.BEISEN,
                "confirmed_jd": self.position.source_jd,
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )
        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        self.assertIn("岗位说明 V1 已确认", response.json()["message"])
        self.assertEqual(response.json()["state"]["code"], "pending_rule")

    def test_rule_list_links_to_read_only_rule_detail(self):
        confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        rule = self.publish_current_rule()
        client = self.authenticated_client(self.hr)

        response = client.get(reverse("analysis:rule_list"))
        self.assertContains(
            response,
            reverse("analysis:rule_detail", args=[rule.pk]),
        )

        response = client.get(reverse("analysis:rule_detail", args=[rule.pk]))
        self.assertContains(response, f"规则 V{rule.version}")
        self.assertContains(response, "评分维度")

    def test_docx_and_xlsx_reference_files_are_parsed(self):
        document = Document()
        table = document.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "岗位名称"
        table.rows[0].cells[1].text = "岗位职责"
        table.rows[0].cells[2].text = "任职要求"
        table.rows[1].cells[0].text = "数据工程师"
        table.rows[1].cells[1].text = "负责数据平台"
        table.rows[1].cells[2].text = "三年以上经验"
        docx_buffer = BytesIO()
        document.save(docx_buffer)
        parsed_jobs = parse_job_summary_docx(docx_buffer.getvalue())
        self.assertEqual(parsed_jobs[0]["title"], "数据工程师")
        self.assertIn("三年以上经验", parsed_jobs[0]["jd"])

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.append(["负责人", "岗位", "负责人邮箱"])
        worksheet.append(["张三", "数据工程师", "zhangsan@example.com"])
        xlsx_buffer = BytesIO()
        workbook.save(xlsx_buffer)
        parsed_reviewers = parse_reviewer_mapping_xlsx(xlsx_buffer.getvalue())
        self.assertEqual(
            parsed_reviewers[0]["metadata"]["reviewers"][0]["email"],
            "zhangsan@example.com",
        )

    def test_publishing_reference_archives_previous_version(self):
        first = self.create_document_position("第一版岗位").reference_document
        first.version = 1
        first.save(update_fields=["version"])
        second = ReferenceDocument.objects.create(
            name="第二版",
            document_type=ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX,
            content_hash="second",
            version=2,
            uploaded_by=self.hr,
        )
        publish_reference_document(second, self.hr)
        first.refresh_from_db()
        second.refresh_from_db()
        self.assertEqual(first.status, ReferenceDocument.Status.ARCHIVED)
        self.assertEqual(second.status, ReferenceDocument.Status.ACTIVE)

    def test_confirmed_document_match_can_add_document_reviewers(self):
        document_position = self.create_document_position("AI工程师")
        mapping_document = ReferenceDocument.objects.create(
            name="负责人表",
            document_type=ReferenceDocument.DocumentType.REVIEWER_MAPPING_XLSX,
            content_hash="reviewers",
            version=1,
            status=ReferenceDocument.Status.ACTIVE,
            uploaded_by=self.hr,
        )
        DocumentPosition.objects.create(
            reference_document=mapping_document,
            title="AI工程师",
            normalized_title="ai工程师",
            metadata={
                "reviewers": [
                    {"name": "算法负责人", "email": "ai-owner@example.com"}
                ]
            },
        )
        self.configuration.document_position = document_position
        self.configuration.match_status = PositionConfiguration.MatchStatus.CONFIRMED
        self.configuration.save()
        self.assertEqual(
            apply_document_reviewers(self.configuration, self.hr),
            1,
        )
        link = self.position.reviewer_links.get()
        self.assertEqual(link.source_type, PositionReviewer.SourceType.DOCUMENT)

    def test_confirm_match_with_dirty_title_and_auto_applies_reviewers(self):
        doc = ReferenceDocument.objects.create(
            name="招聘汇总",
            document_type=ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX,
            content_hash="summary-dirty",
            version=2,
            status=ReferenceDocument.Status.ACTIVE,
            uploaded_by=self.hr,
        )
        doc_pos = DocumentPosition.objects.create(
            reference_document=doc,
            title="3、TK运营/Tiktok运营  招聘负责人：吴晨静、苏碧龙",
            normalized_title="3tk运营tiktok运营招聘负责人吴晨静苏碧龙",
            aliases=["TK运营", "Tiktok运营"],
            jd="负责TK店铺运营",
        )
        mapping_doc = ReferenceDocument.objects.create(
            name="负责人表",
            document_type=ReferenceDocument.DocumentType.REVIEWER_MAPPING_XLSX,
            content_hash="reviewers-dirty",
            version=2,
            status=ReferenceDocument.Status.ACTIVE,
            uploaded_by=self.hr,
        )
        DocumentPosition.objects.create(
            reference_document=mapping_doc,
            title="Tiktok运营",
            normalized_title="tiktok运营",
            aliases=["TK运营/Tiktok运营"],
            metadata={
                "reviewers": [
                    {"name": "吴晨静", "email": "wuchenjing@nuptio.net"},
                    {"name": "苏碧龙", "email": "subilong@nuptio.net"},
                ]
            },
        )
        self.configuration.document_position = doc_pos
        self.configuration.match_status = PositionConfiguration.MatchStatus.CONFIRMED
        self.configuration.save()

        count = apply_document_reviewers(self.configuration, self.hr)
        self.assertEqual(count, 2)
        links = self.position.reviewer_links.all()
        self.assertEqual(links.count(), 2)
        self.assertEqual(
            set(links.values_list("reviewer__name", flat=True)),
            {"吴晨静", "苏碧龙"},
        )

    def test_review_is_rejected_when_position_configuration_is_incomplete(self):
        application = self.create_application(
            applicant_id="REVIEW-CONFIG",
            position=self.position,
        )
        reviewer = Reviewer.objects.create(
            name="审核人",
            email="review-config@example.com",
        )
        with self.assertRaisesMessage(ReviewError, "确认岗位说明"):
            create_review_batch(
                self.position,
                [application.pk],
                reviewer,
                self.hr,
            )

    def test_historical_position_is_never_runnable(self):
        self.position.status = Position.Status.HISTORICAL
        self.position.save(update_fields=["status"])
        state = configuration_state(self.position)
        self.assertEqual(state.code, "historical")
        self.assertFalse(state.can_run)

    def test_no_document_match_can_still_confirm_beisen_content(self):
        self.configuration.match_status = PositionConfiguration.MatchStatus.NO_MATCH
        self.configuration.save(update_fields=["match_status"])
        decision = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        self.assertEqual(decision.confirmed_jd, self.position.source_jd)

    def test_ai_diff_failure_does_not_block_manual_confirmation(self):
        client = self.authenticated_client(self.hr)
        response = client.post(
            reverse("recruitment:configuration_ai_diff", args=[self.position.pk])
        )
        self.assertRedirects(
            response,
            reverse("recruitment:configuration_detail", args=[self.position.pk]),
        )
        response = client.post(
            reverse("recruitment:configuration_confirm_jd", args=[self.position.pk]),
            {
                "decision_type": PositionJdDecision.DecisionType.BEISEN,
                "confirmed_jd": self.position.source_jd,
            },
        )
        self.assertRedirects(
            response,
            reverse("recruitment:configuration_detail", args=[self.position.pk]),
        )
        self.assertTrue(self.position.jd_decisions.filter(is_current=True).exists())

    def test_hr_can_upload_and_publish_reference_file(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "岗位名称"
        table.rows[0].cells[1].text = "岗位说明"
        table.rows[1].cells[0].text = "测试岗位"
        table.rows[1].cells[1].text = "负责测试"
        buffer = BytesIO()
        document.save(buffer)
        upload = SimpleUploadedFile(
            "岗位汇总.docx",
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        client = self.authenticated_client(self.hr)
        response = client.post(
            reverse("recruitment:reference_documents"),
            {
                "document_type": ReferenceDocument.DocumentType.JOB_SUMMARY_DOCX,
                "file": upload,
            },
        )
        self.assertRedirects(response, reverse("recruitment:reference_documents"))
        reference = ReferenceDocument.objects.latest("pk")
        self.assertEqual(reference.positions.count(), 1)
        response = client.post(
            reverse("recruitment:reference_publish", args=[reference.pk])
        )
        self.assertRedirects(response, reverse("recruitment:reference_documents"))
        reference.refresh_from_db()
        self.assertEqual(reference.status, ReferenceDocument.Status.ACTIVE)

    def test_rule_generation_connection_failure_shows_user_message(self):
        confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )

        class FailingCompletions:
            def create(self, **kwargs):
                raise APIConnectionError(
                    request=httpx.Request("POST", "https://model.example/v1")
                )

        class FailingClient:
            class Chat:
                completions = FailingCompletions()

            chat = Chat()

        client = self.authenticated_client(self.hr)
        with patch(
            "analysis.integrations.model.OpenAI",
            return_value=FailingClient(),
        ):
            response = client.post(
                reverse("analysis:rule_generate", args=[self.position.pk]),
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "暂时无法连接智能分析服务")
        self.assertNotContains(response, "APIConnectionError")

    def test_rule_generation_accepts_model_score_ranges(self):
        confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )

        class RangeRuleGateway:
            def analyze(self, system_prompt, user_prompt):
                return {
                    "payload": {
                        "evaluation_jd": self.position.source_jd,
                        "hard_requirements": [{"name": "模型研发经验"}],
                        "dimensions": [
                            {"name": "相关经验", "weight": "60"},
                            {"name": "工程能力", "weight": "40"},
                        ],
                        "bonus_items": [{"name": "模型上线经验"}],
                        "rating_thresholds": {
                            "priority": [80, 100],
                            "review": [60, 79],
                        },
                    },
                    "input_tokens": 100,
                    "output_tokens": 100,
                }

        gateway = RangeRuleGateway()
        gateway.position = self.position
        client = self.authenticated_client(self.hr)
        with patch(
            "analysis.services.rules.ModelGateway",
            return_value=gateway,
        ):
            response = client.post(
                reverse("analysis:rule_generate", args=[self.position.pk]),
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        rule = self.position.rule_versions.get()
        self.assertEqual(rule.rating_thresholds["priority"], 80)
        self.assertEqual(rule.rating_thresholds["review"], 60)
        self.assertNotContains(response, "TypeError")

    def create_rule_draft(self):
        decision = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        return PositionRuleVersion.objects.create(
            position=self.position,
            version=1,
            evaluation_jd=decision.confirmed_jd,
            source_jd_snapshot=decision.confirmed_jd,
            jd_decision=decision,
            hard_requirements=[
                {
                    "name": "学历与专业",
                    "description": "本科及以上学历，计算机相关专业。",
                }
            ],
            dimensions=[
                {
                    "name": "算法能力",
                    "weight": 100,
                    "description": "评估算法基础和项目应用能力。",
                }
            ],
            bonus_items=[
                {
                    "name": "大数据经验",
                    "description": "熟悉常见大数据处理平台。",
                }
            ],
            rating_thresholds={"priority": 80, "review": 60},
            created_by=self.hr,
        )

    def test_rule_edit_page_uses_readable_editors_for_hr_and_admin(self):
        rule = self.create_rule_draft()
        url = reverse(
            "analysis:rule_edit_version",
            args=[self.position.pk, rule.pk],
        )

        for user in (self.hr, self.admin):
            with self.subTest(role=user.role):
                response = self.authenticated_client(user).get(url)
                self.assertEqual(response.status_code, 200)
                self.assertContains(response, "规则草稿")
                self.assertContains(response, "要求名称")
                self.assertContains(response, "维度名称")
                self.assertContains(response, "评估说明")
                self.assertContains(response, "优先推荐起始分")
                self.assertContains(response, 'data-rule-editor')
                self.assertNotContains(response, "默认包含 priority/review/low")
                self.assertContains(response, "取消等待")

    def test_rule_edit_saves_readable_form_values_as_structured_rule(self):
        rule = self.create_rule_draft()
        response = self.authenticated_client(self.hr).post(
            reverse(
                "analysis:rule_edit_version",
                args=[self.position.pk, rule.pk],
            ),
            {
                "hard_requirements": json.dumps(
                    [
                        {
                            "name": "相关经验",
                            "description": "至少两年算法项目经验。",
                        }
                    ]
                ),
                "dimensions": json.dumps(
                    [
                        {
                            "name": "算法与模型能力",
                            "weight": 70,
                            "description": "评估算法选型和模型优化能力。",
                        },
                        {
                            "name": "工程实现",
                            "weight": 30,
                            "description": "评估代码质量和交付能力。",
                        },
                    ]
                ),
                "bonus_items": json.dumps(
                    [
                        {
                            "name": "模型上线经验",
                            "description": "有生产环境部署经验。",
                        }
                    ]
                ),
                "priority_threshold": "85",
                "review_threshold": "65",
            },
        )

        self.assertRedirects(
            response,
            reverse(
                "analysis:rule_edit_version",
                args=[self.position.pk, rule.pk],
            ),
        )
        rule.refresh_from_db()
        self.assertEqual(rule.hard_requirements[0]["name"], "相关经验")
        self.assertEqual(rule.dimensions[0]["weight"], 70)
        self.assertEqual(rule.dimensions[1]["weight"], 30)
        self.assertEqual(rule.bonus_items[0]["name"], "模型上线经验")
        self.assertEqual(
            rule.rating_thresholds,
            {"priority": 85, "review": 65},
        )

    def test_rule_edit_shows_readable_validation_messages(self):
        rule = self.create_rule_draft()
        response = self.authenticated_client(self.hr).post(
            reverse(
                "analysis:rule_edit_version",
                args=[self.position.pk, rule.pk],
            ),
            {
                "hard_requirements": "[]",
                "dimensions": json.dumps(
                    [
                        {
                            "name": "算法能力",
                            "weight": 60,
                            "description": "评估算法基础。",
                        }
                    ]
                ),
                "bonus_items": "[]",
                "priority_threshold": "60",
                "review_threshold": "70",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "所有评分维度的权重合计必须为 100")
        self.assertContains(response, "优先推荐起始分必须高于建议复核起始分")
        rule.refresh_from_db()
        self.assertEqual(rule.dimensions[0]["weight"], 100)

    @override_settings(
        MODEL_API_KEY="test-key",
        MODEL_BASE_URL="https://model.example/v1",
        MODEL_NAME="test-model",
        MODEL_REQUEST_TIMEOUT=45,
    )
    def test_model_gateway_uses_bounded_request_timeout(self):
        response = SimpleNamespace(
            choices=[
                SimpleNamespace(
                    message=SimpleNamespace(content='{"status": "ok"}')
                )
            ],
            usage=None,
        )
        client = SimpleNamespace(
            chat=SimpleNamespace(
                completions=SimpleNamespace(
                    create=lambda **kwargs: response,
                )
            )
        )

        with patch(
            "analysis.integrations.model.OpenAI",
            return_value=client,
        ) as openai_client:
            from analysis.integrations.model import ModelGateway

            ModelGateway().analyze("system", "user")

        openai_client.assert_called_once_with(
            api_key="test-key",
            base_url="https://model.example/v1",
            timeout=45,
            max_retries=0,
        )

    def test_task_overlay_script_can_cancel_and_restore_page(self):
        script = (Path(settings.BASE_DIR) / "static" / "js" / "app.js").read_text(
            encoding="utf-8"
        )

        self.assertIn("window.stop()", script)
        self.assertIn('window.addEventListener("pageshow"', script)
        self.assertIn("已停止页面等待", script)

    def test_rule_generation_cancel_endpoint_marks_operation(self):
        operation = RuleGenerationOperation.objects.create(
            position=self.position,
            requested_by=self.hr,
        )

        response = self.authenticated_client(self.hr).post(
            reverse(
                "analysis:rule_generation_cancel",
                args=[self.position.pk],
            ),
            {"operation_id": str(operation.pk)},
        )

        self.assertEqual(response.status_code, 200)
        self.assertJSONEqual(response.content, {"ok": True})
        operation.refresh_from_db()
        self.assertEqual(
            operation.status,
            RuleGenerationOperation.Status.CANCELLATION_REQUESTED,
        )

    def test_cancelled_rule_generation_does_not_call_model_or_save_draft(self):
        confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        operation_id = uuid4()
        client = self.authenticated_client(self.hr)
        client.post(
            reverse(
                "analysis:rule_generation_cancel",
                args=[self.position.pk],
            ),
            {"operation_id": str(operation_id)},
        )

        with patch("analysis.services.rules.ModelGateway") as model_gateway:
            response = client.post(
                reverse("analysis:rule_generate", args=[self.position.pk]),
                {"operation_id": str(operation_id)},
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "岗位规则草稿生成已取消")
        model_gateway.assert_not_called()
        self.assertFalse(self.position.rule_versions.exists())
        operation = RuleGenerationOperation.objects.get(pk=operation_id)
        self.assertEqual(
            operation.status,
            RuleGenerationOperation.Status.CANCELLED,
        )

    def test_cancelling_during_model_request_discards_generated_rule(self):
        confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
        )
        operation_id = uuid4()

        class CancellingGateway:
            def analyze(self, system_prompt, user_prompt):
                RuleGenerationOperation.objects.filter(pk=operation_id).update(
                    status=RuleGenerationOperation.Status.CANCELLATION_REQUESTED
                )
                return {
                    "payload": {
                        "evaluation_jd": self.position.source_jd,
                        "hard_requirements": [{"name": "模型研发经验"}],
                        "dimensions": [{"name": "相关经验", "weight": 100}],
                        "bonus_items": [],
                        "rating_thresholds": {"priority": 80, "review": 60},
                    },
                    "input_tokens": 100,
                    "output_tokens": 100,
                }

        gateway = CancellingGateway()
        gateway.position = self.position
        with patch(
            "analysis.services.rules.ModelGateway",
            return_value=gateway,
        ):
            response = self.authenticated_client(self.hr).post(
                reverse("analysis:rule_generate", args=[self.position.pk]),
                {"operation_id": str(operation_id)},
                follow=True,
            )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "岗位规则草稿生成已取消")
        self.assertFalse(self.position.rule_versions.exists())
        operation = RuleGenerationOperation.objects.get(pk=operation_id)
        self.assertEqual(
            operation.status,
            RuleGenerationOperation.Status.CANCELLED,
        )

    def test_pinyin_name_to_email_and_api(self):
        from recruitment.services.pinyin import name_to_pinyin, name_to_reviewer_email
        from recruitment.forms_configuration import ReviewerForm

        self.assertEqual(name_to_pinyin("张三"), "zhangsan")
        self.assertEqual(name_to_pinyin("张玉凡"), "zhangyufan")
        self.assertEqual(name_to_reviewer_email("张三"), "zhangsan@nuptio.net")
        self.assertEqual(name_to_reviewer_email("李四"), "lisi@nuptio.net")

        # Test ReviewerForm auto-generation
        form = ReviewerForm({"name": "张三", "email": ""})
        self.assertTrue(form.is_valid())
        self.assertEqual(form.cleaned_data["email"], "zhangsan@nuptio.net")

        # Test ReviewerForm with custom email
        form_custom = ReviewerForm({"name": "张三", "email": "custom@example.com"})
        self.assertTrue(form_custom.is_valid())
        self.assertEqual(form_custom.cleaned_data["email"], "custom@example.com")

        # Test API endpoint
        client = self.authenticated_client(self.hr)
        res = client.get(reverse("recruitment:pinyin_email_api"), {"name": "张三"})
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.json(), {"ok": True, "name": "张三", "email": "zhangsan@nuptio.net"})

        # Test adding reviewer via view without email
        add_res = client.post(
            reverse("recruitment:configuration_add_reviewer", args=[self.position.pk]),
            {"name": "张三", "email": ""},
            follow=True,
        )
        self.assertEqual(add_res.status_code, 200)
        self.assertContains(add_res, "审核负责人已保存")
        link = self.position.reviewer_links.get(reviewer__name="张三")
        self.assertEqual(link.reviewer.email, "zhangsan@nuptio.net")

    def test_build_merged_jd_and_configuration_page_payload(self):
        from recruitment.services.configuration import build_merged_jd

        beisen = (
            "岗位职责\n"
            "1. 负责市场调研，分析热卖产品\n"
            "2. 负责竞品分析\n"
            "任职要求\n"
            "1. 本科以上学历，英语6级\n"
            "2. 工作满3年以上"
        )
        doc = (
            "工作职责：\n"
            "1. 负责市场调研，分析热卖产品\n"
            "2. 负责供应商对接与谈判\n"
            "任职资格：\n"
            "1. 本科以上学历，英语6级\n"
            "2. 具有家居产品开发经验者优先"
        )
        merged = build_merged_jd(beisen, doc)
        self.assertIn("岗位职责：", merged)
        self.assertIn("负责市场调研，分析热卖产品", merged)
        self.assertIn("负责竞品分析", merged)
        self.assertIn("负责供应商对接与谈判", merged)
        self.assertIn("任职要求：", merged)
        self.assertIn("本科以上学历，英语6级", merged)
        self.assertIn("工作满3年以上", merged)
        self.assertIn("具有家居产品开发经验者优先", merged)

        # Test configuration detail view contains merged_jd data script
        confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        client = self.authenticated_client(self.hr)
        response = client.get(
            reverse("recruitment:configuration_detail", args=[self.position.pk])
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="beisen-jd-data"')
        self.assertContains(response, 'id="merged-jd-data"')
        self.assertContains(response, 'id="jd-rules-map-data"')
        self.assertContains(response, "data-jd-decision-form")
        self.assertContains(response, "jd-draft-notice")
        self.assertContains(response, 'id="jd-history-modal"')
        self.assertContains(response, "data-apply-jd-id")

    def test_clean_diff_summary_text_and_plain_format(self):
        from analysis.services.jd_comparison import clean_diff_summary_text

        raw_md = (
            "### 1. 核心职责差异\n"
            "- 北森强调**日常运维**，参考资料偏向**架构设计**。\n\n"
            "### 2. 硬性要求差异\n"
            "- 学历要求一致，工作年限`3年以上`。"
        )
        cleaned = clean_diff_summary_text(raw_md)
        self.assertNotIn("###", cleaned)
        self.assertNotIn("**", cleaned)
        self.assertNotIn("`", cleaned)
        self.assertIn("【核心职责差异】", cleaned)
        self.assertIn("【硬性要求差异】", cleaned)

    def test_confirm_jd_auto_adopts_matching_rule(self):
        # 1. Confirm V1 JD and create Rule V1, then publish Rule V1
        d1 = confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        r1 = PositionRuleVersion.objects.create(
            position=self.position,
            version=1,
            jd_decision=d1,
            evaluation_jd=d1.confirmed_jd,
            source_jd_snapshot=d1.confirmed_jd,
            status=PositionRuleVersion.Status.DRAFT,
            created_by=self.hr,
        )
        r1.publish(self.hr)
        self.assertEqual(r1.status, PositionRuleVersion.Status.PUBLISHED)

        # 2. Confirm V2 JD (new text)
        d2 = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.MANUAL,
            self.hr,
            confirmed_jd="全新版岗位说明内容",
        )
        r1.refresh_from_db()
        # Rule V1 should be archived because V2 has no rule yet
        self.assertEqual(r1.status, PositionRuleVersion.Status.ARCHIVED)

        # 3. Create Rule V2 under V2 JD and publish it
        r2 = PositionRuleVersion.objects.create(
            position=self.position,
            version=2,
            jd_decision=d2,
            evaluation_jd=d2.confirmed_jd,
            source_jd_snapshot=d2.confirmed_jd,
            status=PositionRuleVersion.Status.DRAFT,
            created_by=self.hr,
        )
        r2.publish(self.hr)
        self.assertEqual(r2.status, PositionRuleVersion.Status.PUBLISHED)

        # 4. Re-confirm V1 JD (using V1 content)
        d1_reactivated = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.BEISEN,
            self.hr,
            confirmed_jd=d1.confirmed_jd,
        )
        self.assertEqual(d1_reactivated.pk, d1.pk)
        self.assertTrue(d1_reactivated.is_current)

        r1.refresh_from_db()
        r2.refresh_from_db()
        # Rule V1 should now be auto-activated to PUBLISHED, and Rule V2 archived!
        self.assertEqual(r1.status, PositionRuleVersion.Status.PUBLISHED)
        self.assertEqual(r2.status, PositionRuleVersion.Status.ARCHIVED)

    def test_jd_decision_deletion_and_protection(self):
        from recruitment.services.configuration import delete_jd_decision

        # Create two JD decisions
        d1 = confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        d2 = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.MANUAL,
            self.hr,
            confirmed_jd="修改后的岗位说明内容",
        )
        d1.refresh_from_db()
        self.assertFalse(d1.is_current)
        self.assertTrue(d2.is_current)

        # 1. Deleting current active JD raises ValueError
        with self.assertRaises(ValueError) as ctx:
            delete_jd_decision(d2, self.hr)
        self.assertIn("当前正在生效", str(ctx.exception))

        # 2. Deleting non-current JD without rules succeeds
        delete_jd_decision(d1, self.hr)
        self.assertFalse(PositionJdDecision.objects.filter(pk=d1.pk).exists())

        # 3. Create another JD and link a rule to it
        d3 = confirm_jd(
            self.position,
            PositionJdDecision.DecisionType.MANUAL,
            self.hr,
            confirmed_jd="第三版岗位说明内容",
        )
        d2.refresh_from_db()
        rule = PositionRuleVersion.objects.create(
            position=self.position,
            version=1,
            jd_decision=d2,
            evaluation_jd=d2.confirmed_jd,
            source_jd_snapshot=d2.confirmed_jd,
            status=PositionRuleVersion.Status.DRAFT,
            created_by=self.hr,
        )

        # Non-admin deleting d2 raises ValueError
        with self.assertRaises(ValueError) as ctx:
            delete_jd_decision(d2, self.hr)
        self.assertIn("无法直接删除", str(ctx.exception))

        # Admin with force=True deletes d2 and unlinks rule
        delete_jd_decision(d2, self.admin, force=True)
        self.assertFalse(PositionJdDecision.objects.filter(pk=d2.pk).exists())
        rule.refresh_from_db()
        self.assertIsNone(rule.jd_decision)

    def test_rule_version_deletion_and_publishing(self):
        from analysis.services.rules import delete_rule_version

        decision = confirm_jd(self.position, PositionJdDecision.DecisionType.BEISEN, self.hr)
        r1 = PositionRuleVersion.objects.create(
            position=self.position,
            version=1,
            jd_decision=decision,
            evaluation_jd=decision.confirmed_jd,
            source_jd_snapshot=decision.confirmed_jd,
            status=PositionRuleVersion.Status.DRAFT,
            created_by=self.hr,
        )
        r1.publish(self.hr)
        self.assertEqual(r1.status, PositionRuleVersion.Status.PUBLISHED)

        # Cannot delete active published rule
        with self.assertRaises(ValueError) as ctx:
            delete_rule_version(r1, self.hr)
        self.assertIn("已发布", str(ctx.exception))

        # Create r2 and publish it (r1 becomes archived)
        r2 = PositionRuleVersion.objects.create(
            position=self.position,
            version=2,
            jd_decision=decision,
            evaluation_jd=decision.confirmed_jd,
            source_jd_snapshot=decision.confirmed_jd,
            status=PositionRuleVersion.Status.DRAFT,
            created_by=self.hr,
        )
        r2.publish(self.hr)
        r1.refresh_from_db()
        self.assertEqual(r1.status, PositionRuleVersion.Status.ARCHIVED)

        # Can delete archived r1 when no analysis items
        delete_rule_version(r1, self.hr)
        self.assertFalse(PositionRuleVersion.objects.filter(pk=r1.pk).exists())

        # Test view endpoint for rule delete
        r3 = PositionRuleVersion.objects.create(
            position=self.position,
            version=3,
            jd_decision=decision,
            evaluation_jd=decision.confirmed_jd,
            source_jd_snapshot=decision.confirmed_jd,
            status=PositionRuleVersion.Status.DRAFT,
            created_by=self.hr,
        )
        client = self.authenticated_client(self.hr)
        res = client.post(reverse("analysis:rule_delete", args=[r3.pk]))
        self.assertEqual(res.status_code, 302)
        self.assertFalse(PositionRuleVersion.objects.filter(pk=r3.pk).exists())


