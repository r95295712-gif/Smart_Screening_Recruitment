from io import BytesIO

from docx import Document


def docx_resume_bytes():
    document = Document()
    document.add_heading("软件工程师简历", 0)
    document.add_paragraph(
        "候选人具有八年软件研发经验，负责过招聘系统、数据同步、异步任务、"
        "数据库设计、系统部署和跨团队协作。" * 12
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class FakeITalentClient:
    def __init__(self):
        self.origin_bytes = docx_resume_bytes()

    def get_positions(self, position_ids):
        return {
            "data": {
                "items": [
                    {
                        "jobId": "P-1",
                        "jobTitle": "后端工程师",
                        "jobType": ["技术"],
                        "status": 1,
                        "duty": "负责招聘系统后端开发。",
                        "require": "熟悉 Python。",
                    },
                    {
                        "jobId": "P-2",
                        "jobTitle": "平台工程师",
                        "jobType": ["技术"],
                        "status": 1,
                        "duty": "负责平台工程建设。",
                    },
                ]
            }
        }

    def iter_applicant_ids(self, start_time, end_time, time_type=2):
        yield ["A-1"]

    def get_profiles(self, applicant_ids):
        return {
            "data": {
                "items": [
                    {
                        "applicantId": "A-1",
                        "fieldValues": [
                            {"name": "Name", "value": "张三", "text": "张三"},
                            {
                                "name": "Mobile",
                                "value": "13800138000",
                                "text": "13800138000",
                            },
                            {
                                "name": "Email",
                                "value": "zhangsan@example.com",
                                "text": "zhangsan@example.com",
                            },
                            {
                                "name": "LastCompany",
                                "value": "示例科技",
                                "text": "示例科技",
                            },
                            {
                                "name": "LastSchool",
                                "value": "示例大学",
                                "text": "示例大学",
                            },
                        ],
                    }
                ]
            }
        }

    def get_resume_module(self, applicant_ids, module_code):
        return {
            "data": {
                "items": [
                    {
                        "applicantId": "A-1",
                        "moduleInfo": [
                            [
                                {
                                    "name": "SkillName",
                                    "value": f"{module_code} 示例内容",
                                    "text": f"{module_code} 示例内容",
                                }
                            ]
                        ],
                    }
                ]
            }
        }

    def get_applications(self, applicant_ids):
        return {
            "data": {
                "items": [
                    {
                        "applicationId": "APP-1",
                        "applicantId": "A-1",
                        "positionId": "P-1",
                        "positionName": "后端工程师",
                        "appliedTime": "2026-08-01T10:00:00+08:00",
                        "source": "官网",
                    },
                    {
                        "applicationId": "APP-2",
                        "applicantId": "A-1",
                        "positionId": "P-2",
                        "positionName": "平台工程师",
                        "appliedTime": "2026-08-02T10:00:00+08:00",
                        "source": "内推",
                    },
                ]
            }
        }

    def get_resume_file_info(self, applicant_id, *, origin=True):
        if origin:
            return {
                "data": {
                    "downloadUrl": "//dfiles.example/origin/A-1.docx?sig=fake",
                    "dfsPath": "/origin/A-1.docx",
                }
            }
        return {
            "data": {
                "downloadUrl": "//dfiles.example/standard/A-1.pdf?sig=fake",
                "dfsPath": "/standard/A-1.pdf",
            }
        }

    def download_file(self, url):
        if "/origin/" in url:
            return (
                self.origin_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        return b"%PDF-1.4\n% fake preview\n", "application/pdf"


class FakeModelGateway:
    def analyze(self, system_prompt, user_prompt):
        return {
            "payload": {
                "score": 86,
                "hard_requirement_results": [
                    {"requirement": "后端经验", "status": "满足", "evidence": "八年研发经验"}
                ],
                "dimension_results": [
                    {"dimension": "相关经验", "score": 40, "evidence": "招聘系统"}
                ],
                "strengths": ["相关系统经验丰富"],
                "risks": ["团队规模信息不足"],
                "missing_information": ["期望薪资"],
                "interview_focus": ["核实系统规模"],
                "interview_questions": ["请介绍最复杂的数据同步问题。"],
            },
            "input_tokens": 1000,
            "output_tokens": 300,
        }


class FakeRuleGateway:
    def analyze(self, system_prompt, user_prompt):
        return {
            "payload": {
                "evaluation_jd": "负责后端系统研发和稳定性建设。",
                "hard_requirements": [{"name": "后端研发经验"}],
                "dimensions": [
                    {"name": "相关经验", "weight": 60},
                    {"name": "系统设计", "weight": 40},
                ],
                "bonus_items": [{"name": "招聘系统经验"}],
                "rating_thresholds": {"priority": 80, "review": 60},
            },
            "input_tokens": 100,
            "output_tokens": 100,
        }


class RecordingRuleGateway(FakeRuleGateway):
    def __init__(self):
        self.user_prompt = ""

    def analyze(self, system_prompt, user_prompt):
        self.user_prompt = user_prompt
        return super().analyze(system_prompt, user_prompt)
